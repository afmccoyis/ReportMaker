import openpyxl
import os
import docx

#print(os.getcwd())

wb = openpyxl.load_workbook('Domestic_FITBook_Ver_5.1.3_14JUN2021.xlsm', data_only=True, keep_vba=True)
doc = docx.Document()
#print(wb.get_sheet_names())
#Tracker_Name = []
#Tracker_Value = []
#y = 0


def WhatsNext(x):
    if x == 'PASS':
        print(f"Hell Yeah")
        doc.add_paragraph(x)
        #doc.save('test.docx')
    elif x == 'FAIL':
        print(f"Hell No")
        doc.add_paragraph(x)
    else:
        print("Ok")

def Info_Page(ExeSumm):
    #Vendor Name
    VendorName = ExeSumm['B3'].value
    #Test Location
    Test_Location = ExeSumm['F3'].value
    #Model Name
    ModelName = ExeSumm['B4'].value
    #Software
    SoftwareName = ExeSumm['F5'].value

#print(os.getcwd())
def Tab_4G_Sum(ExeSumm):
    # Marginal VCP
    Marg_MO = ExeSumm['C9'].value
    if Marg_MO == 'PASS':
        doc.add_paragraph("Marginal Voice MO = Pass")
    elif Marg_MO == 'FAIL':
        doc.add_paragraph("Marginal Voice MO = Failed")
    #WhatsNext(Marg_MO)
    Marg_MT = ExeSumm['C10'].value
    if Marg_MT == 'PASS':
        doc.add_paragraph("Marginal Voice MT = Pass")
    elif Marg_MT == 'FAIL':
        doc.add_paragraph("Marginal Voice MT = Failed")
    #WhatsNext(Marg_MT)
    Marg_LC = ExeSumm['C11'].value
    if Marg_LC == 'PASS':
        doc.add_paragraph("Marginal Voice LC = Pass")
    elif Marg_LC == 'FAIL':
        doc.add_paragraph("Marginal Voice LC = Failed")
    #WhatsNext(Marg_LC)
    #doc.save('test.docx')
    # Mixed VCP
    Mixed_MO = ExeSumm['C15'].value
    if Mixed_MO == 'PASS':
        doc.add_paragraph("Mixed Voice MO = Pass")
    elif Mixed_MO == 'FAIL':
        doc.add_paragraph("Mixed Voice MO = Failed")
    Mixed_MT = ExeSumm['C16'].value
    if Mixed_MT == 'PASS':
        doc.add_paragraph("Mixed Voice MT = Pass")
    elif Mixed_MT == 'FAIL':
        doc.add_paragraph("Mixed Voice MT = Failed")
    Mixed_LC = ExeSumm['C17'].value
    if Mixed_LC == 'PASS':
        doc.add_paragraph("Mixed Voice LC = Pass")
    elif Mixed_LC == 'FAIL':
        doc.add_paragraph("Mixed Voice LC = Failed")

    # Data Mobility
    Ping_MO = ExeSumm['C21'].value
    if Ping_MO == 'PASS':
        doc.add_paragraph("3.3.9 LTE Data MO = Pass")
    elif Ping_MO == 'FAIL':
        doc.add_paragraph("3.3.9 LTE Data MO = Failed")
    Data_LC = ExeSumm['C23'].value
    if Data_LC == 'PASS':
        doc.add_paragraph("3.3.11 LTE Data LC = Pass")
    elif Ping_MO == 'FAIL':
        doc.add_paragraph("3.3.11 LTE Data LC = Failed")

    # LTE Data Performance
    B13_USB_DL = ExeSumm['G9'].value
    if B13_USB_DL == 'PASS':
        doc.add_paragraph("3.3.1 B13 LTE USB DL = Pass")
    elif B13_USB_DL == 'FAIL':
        doc.add_paragraph("3.3.1 B13 LTE USB DL = Failed")
    B13_USB_UL = ExeSumm['G10'].value
    if B13_USB_UL == 'PASS':
        doc.add_paragraph("3.3.2 B13 LTE USB UL = Pass")
    elif B13_USB_UL == 'FAIL':
        doc.add_paragraph("3.3.2 B13 LTE USB UL = Failed")
    B13_MHS_DL = ExeSumm['G11'].value
    if B13_MHS_DL == 'PASS':
        doc.add_paragraph("3.3.3 B13 LTE MHS DL = Pass")
    elif B13_MHS_DL == 'FAIL':
        doc.add_paragraph("3.3.3 B13 LTE MHS DL = Failed")
    B13_MHS_UL = ExeSumm['G12'].value
    if B13_MHS_UL == 'PASS':
        doc.add_paragraph("3.3.4 B13 LTE MHS UL = Pass")
    elif B13_MHS_UL == 'FAIL':
        doc.add_paragraph("3.3.4 B13 LTE MHS UL = Failed")
    B04_USB_DL = ExeSumm['H9'].value
    if B04_USB_DL == 'PASS':
        doc.add_paragraph("3.3.1 B4 LTE USB DL = Pass")
    elif B04_USB_DL == 'FAIL':
        doc.add_paragraph("3.3.1 B4 LTE USB DL = Failed")
    B04_USB_UL = ExeSumm['H10'].value
    if B04_USB_UL == 'PASS':
        doc.add_paragraph("3.3.2 B4 LTE USB UL = Pass")
    elif B04_USB_UL == 'FAIL':
        doc.add_paragraph("3.3.2 B4 LTE USB UL = Failed")
    B04_MHS_DL = ExeSumm['H11'].value
    if B04_MHS_DL == 'PASS':
        doc.add_paragraph("3.3.3 B13 LTE MHS DL = Pass")
    elif B04_MHS_DL == 'FAIL':
        doc.add_paragraph("3.3.3 B13 LTE MHS DL = Failed")
    B04_MHS_UL = ExeSumm['H12'].value
    if B04_MHS_UL == 'PASS':
        doc.add_paragraph("3.3.4 B13 LTE MHS UL = Pass")
    elif B04_MHS_UL == 'FAIL':
        doc.add_paragraph("3.3.4 B13 LTE MHS UL = Failed")
    B02_USB_DL = ExeSumm['I9'].value
    if B02_USB_DL == 'PASS':
        doc.add_paragraph("3.3.1 B2 LTE USB DL = Pass")
    elif B02_USB_DL == 'FAIL':
        doc.add_paragraph("3.3.1 B2 LTE USB DL = Failed")
    B02_USB_UL = ExeSumm['I10'].value
    if B02_USB_UL == 'PASS':
        doc.add_paragraph("3.3.2 B2 LTE USB UL = Pass")
    elif B02_USB_UL == 'FAIL':
        doc.add_paragraph("3.3.2 B2 LTE USB UL = Failed")
    B02_MHS_DL = ExeSumm['I11'].value
    if B02_MHS_DL == 'PASS':
        doc.add_paragraph("3.3.3 B2 LTE MHS DL = Pass")
    elif B02_MHS_DL == 'FAIL':
        doc.add_paragraph("3.3.3 B2 LTE MHS DL = Failed")
    B02_MHS_UL = ExeSumm['I12'].value
    if B02_MHS_UL == 'PASS':
        doc.add_paragraph("3.3.4 B2 LTE MHS UL = Pass")
    elif B02_MHS_UL == 'FAIL':
        doc.add_paragraph("3.3.4 B2 LTE MHS UL = Failed")

    # Video VCP
    Video_MO = ExeSumm['H28'].value
    if Video_MO == 'PASS':
        doc.add_paragraph("Mixed Video MO = Pass")
    elif Video_MO == 'FAIL':
        doc.add_paragraph("Mixed Video MO = Failed")
    Video_MT = ExeSumm['H29'].value
    if Video_MT == 'PASS':
        doc.add_paragraph("Mixed Video MT = Pass")
    elif Video_MT == 'FAIL':
        doc.add_paragraph("Mixed Video MT = Failed")
    Video_LC = ExeSumm['H30'].value
    if Video_LC == 'PASS':
        doc.add_paragraph("Mixed Video LC = Pass")
    elif Video_LC == 'FAIL':
        doc.add_paragraph("Mixed Video LC = Failed")

    # Intermarket VCP
    Int_Video_MO = ExeSumm['J28'].value
    if Int_Video_MO == 'PASS':
        doc.add_paragraph("Intermarket Video MO = Pass")
    elif Int_Video_MO == 'FAIL':
        doc.add_paragraph("Intermarket Video MO = Failed")
    Int_Video_MT = ExeSumm['J29'].value
    if Int_Video_MT == 'PASS':
        doc.add_paragraph("Intermarket Video MT = Pass")
    elif Int_Video_MT == 'FAIL':
        doc.add_paragraph("Intermarket Video MT = Failed")
    Int_Voice_MO = ExeSumm['J31'].value
    if Int_Voice_MO == 'PASS':
        doc.add_paragraph("Intermarket Voice MO = Pass")
    elif Int_Voice_MO == 'FAIL':
        doc.add_paragraph("Intermarket Voice MO = Failed")
    Int_Voice_MT = ExeSumm['J32'].value
    if Int_Voice_MT == 'PASS':
        doc.add_paragraph("Intermarket Voice MT = Pass")
    elif Int_Voice_MT == 'FAIL':
        doc.add_paragraph("Intermarket Voice MT = Failed")

    #4G Legacy Features
    Provision4G = ExeSumm['C26'].value
    Sel_Reg4G = ExeSumm['C27'].value
    Call_Feat = ExeSumm['C28'].value
    Sms = ExeSumm['C29'].value
    Data_Serv = ExeSumm['C30'].value
    Concurrent_Ser4G = ExeSumm['C31'].value
    Domestic_Roam = ExeSumm['C32'].value
    if Provision4G == 'Fail':
        doc.add_paragraph("4G Legacy Features: Service Provisioning = Failed")
    elif Provision4G == 'Pass':
        doc.add_paragraph("4G Legacy Features: Service Provisioning = Pass")
    if Sel_Reg4G == 'Fail':
        doc.add_paragraph("4G Legacy Features: System Sel. And Reg. = Failed")
    elif Sel_Reg4G == 'Pass':
        doc.add_paragraph("4G Legacy Features: System Sel. And Reg. = Pass")
    if Call_Feat == 'Fail':
        doc.add_paragraph("4G Legacy Features: Sub. Call Feats = Failed")
    elif Call_Feat == 'Pass':
        doc.add_paragraph("4G Legacy Features: Sub. Call Feats = Pass")
    if Sms == 'Fail':
        doc.add_paragraph("4G Legacy Features: SMS = Failed")
    elif Sms == 'Pass':
        doc.add_paragraph("4G Legacy Features: SMS = Pass")
    if Data_Serv == 'Fail':
        doc.add_paragraph("4G Legacy Features: Data Service Func = Failed")
    elif Data_Serv == 'Pass':
        doc.add_paragraph("4G Legacy Features: Data Service Func = Pass")
    if Concurrent_Ser4G == 'Fail':
        doc.add_paragraph("4G Legacy Features: Concurrent Services = Failed")
    elif Concurrent_Ser4G == 'Pass':
        doc.add_paragraph("4G Legacy Features: Concurrent Services = Pass")
    if Domestic_Roam == 'Fail':
        doc.add_paragraph("Domestic Roaming = Failed")
    elif Domestic_Roam == 'Pass':
        doc.add_paragraph("Domestic Roaming = Pass")

    #4G VoLTE Features
    InitSetup = ExeSumm['H20'].value
    VvCallEst = ExeSumm['H21'].value
    SuppFea = ExeSumm['H23'].value
    InterTestCase = ExeSumm['H24'].value
    if InitSetup == 'Fail':
        doc.add_paragraph("4G VoLTE Features: Initial Setup = Failed")
    elif InitSetup == 'Pass':
        doc.add_paragraph("4G VoLTE Features: Initial Setup = Pass")
    if VvCallEst == 'Fail':
        doc.add_paragraph("4G VoLTE Features: Voice/Video Call Establishment = Failed")
    elif VvCallEst == 'Pass':
        doc.add_paragraph("4G VoLTE Features: Voice/Video Call Establishment = Pass")
    if SuppFea == 'Fail':
        doc.add_paragraph("4G VoLTE Features: Supplemental Features = Failed")
    elif SuppFea == 'Pass':
        doc.add_paragraph("4G VoLTE Features: Supplemental Features = Pass")
    if InterTestCase == 'Fail':
        doc.add_paragraph("4G VoLTE Features: Interaction Test Cases = Failed")
    elif InterTestCase == 'Pass':
        doc.add_paragraph("4G VoLTE Features: Interaction Test Cases = Pass")
    #doc.save('test.docx')

def Tab_5G_Sum(Exe5GSumm):
    #5G Features
    Provision5G_FR2 = Exe5GSumm['B9'].value
    Sel_Reg5G_FR2 = Exe5GSumm['B10'].value
    Concurrent_Ser5G_FR2 = Exe5GSumm['B11'].value
    Streaming5G_FR2 = Exe5GSumm['B12'].value
    Interoperability5G_FR2 = Exe5GSumm['B13'].value
    Provision5G_FR1 = Exe5GSumm['C9'].value
    Sel_Reg5G_FR1 = Exe5GSumm['C10'].value
    Concurrent_Ser5G_FR1 = Exe5GSumm['C11'].value
    Streaming5G_FR1 = Exe5GSumm['C12'].value
    Interoperability5G_FR1 = Exe5GSumm['C13'].value
    Dss_sec6_5G_FR1 = Exe5GSumm['C14'].value
    if Provision5G_FR2 == 'Fail':
        doc.add_paragraph("5G FR2 Features: Service Provisioning = Failed")
    elif Provision5G_FR2 == 'Pass':
        doc.add_paragraph("5G FR2 Features: Service Provisioning = Pass")
    if Sel_Reg5G_FR2 == 'Fail':
        doc.add_paragraph("5G FR2 Features: System Sel. And Reg. = Failed")
    elif Sel_Reg5G_FR2 == 'Pass':
        doc.add_paragraph("5G FR2 Features: System Sel. And Reg. = Pass")
    if Concurrent_Ser5G_FR2 == 'Fail':
        doc.add_paragraph("5G FR2 Features: Concurrent Services = Failed")
    elif Concurrent_Ser5G_FR2 == 'Pass':
        doc.add_paragraph("5G FR2 Features: Concurrent Services = Pass")
    if Streaming5G_FR2 == 'Fail':
        doc.add_paragraph("5G FR2 Features: Streaming = Failed")
    elif Streaming5G_FR2 == 'Pass':
        doc.add_paragraph("5G FR2 Features: Streaming = Pass")
    if Interoperability5G_FR2 == 'Fail':
        doc.add_paragraph("5G FR2 Features: Interoperability = Failed")
    elif Interoperability5G_FR2 == 'Pass':
        doc.add_paragraph("5G FR2 Features: Interoperability = Pass")

    if Provision5G_FR1 == 'Fail':
        doc.add_paragraph("5G FR1 Features: Service Provisioning = Failed")
    elif Provision5G_FR1 == 'Pass':
        doc.add_paragraph("5G FR1 Features: Service Provisioning = Pass")
    if Sel_Reg5G_FR1 == 'Fail':
        doc.add_paragraph("5G FR1 Features: System Sel. And Reg. = Failed")
    elif Sel_Reg5G_FR1 == 'Pass':
        doc.add_paragraph("5G FR1 Features: System Sel. And Reg. = Pass")
    if Concurrent_Ser5G_FR1 == 'Fail':
        doc.add_paragraph("5G FR1 Features: Concurrent Services = Failed")
    elif Concurrent_Ser5G_FR1 == 'Pass':
        doc.add_paragraph("5G FR1 Features: Concurrent Services = Pass")
    if Streaming5G_FR1 == 'Fail':
        doc.add_paragraph("5G FR1 Features: Streaming = Failed")
    elif Streaming5G_FR1 == 'Pass':
        doc.add_paragraph("5G FR1 Features: Streaming = Pass")
    if Interoperability5G_FR1 == 'Fail':
        doc.add_paragraph("5G FR1 Features: Interoperability = Failed")
    elif Interoperability5G_FR1 == 'Pass':
        doc.add_paragraph("5G FR1 Features: Interoperability = Pass")
    if Dss_sec6_5G_FR1 == 'Fail':
        doc.add_paragraph("DSS Section 6 Testing = Failed")
    elif Dss_sec6_5G_FR1 == 'Pass':
        doc.add_paragraph("DSS Section 6 Testing = Pass")

    #5G Data Performance
    Dss_5G_DL = Exe5GSumm['C18'].value
    Dss_5G_UL = Exe5GSumm['C19'].value
    Fr2_5G_DL = Exe5GSumm['H18'].value
    Fr2_5G_UL = Exe5GSumm['H19'].value
    Fr2_5G_bidir = Exe5GSumm['H20'].value
    if Dss_5G_DL == 'Fail':
        doc.add_paragraph("3.1 5G FR1 Downloads = Failed")
    elif Dss_5G_DL == 'Pass':
        doc.add_paragraph("3.1 5G FR1 Downloads = Pass")
    if Dss_5G_UL == 'Fail':
        doc.add_paragraph("3.2 5G FR1 Uploads = Failed")
    elif Dss_5G_UL == 'Pass':
        doc.add_paragraph("3.2 5G FR1 Uploads = Pass")
    if Fr2_5G_DL == 'Fail':
        doc.add_paragraph("3.1 5G FR2 Downloads = Failed")
    elif Fr2_5G_DL == 'Pass':
        doc.add_paragraph("3.1 5G FR2 Downloads = Pass")
    if Fr2_5G_UL == 'Fail':
        doc.add_paragraph("3.2 5G FR2 Uploads = Failed")
    elif Fr2_5G_UL == 'Pass':
        doc.add_paragraph("3.2 5G FR2 Uploads = Pass")
    if Fr2_5G_bidir == 'Fail':
        doc.add_paragraph("3.3 5G FR2 Bi-directional = Failed")
    elif Fr2_5G_bidir == 'Pass':
        doc.add_paragraph("3.3 5G FR2 Bi-directional = Pass")

    #5G Video Performance
    Voice_Orig5G = Exe5GSumm['H9'].value
    Voice_Term5G = Exe5GSumm['H10'].value
    #Voice_Lc5G = Exe5GSumm['H11'].value
    Video_Orig5G = Exe5GSumm['H12'].value
    Video_Term5G = Exe5GSumm['H13'].value
    Video_Lc5G = Exe5GSumm['H14'].value
    if Voice_Orig5G == 'Fail':
        doc.add_paragraph("4.1 5G Voice MO = Failed")
    elif Voice_Orig5G == 'Pass':
        doc.add_paragraph("4.1 5G Voice MO = Pass")
    if Voice_Term5G == 'Fail':
        doc.add_paragraph("4.1 5G Voice MT = Failed")
    elif Voice_Term5G == 'Pass':
        doc.add_paragraph("4.1 5G Voice MT = Pass")
    if Video_Orig5G == 'Fail':
        doc.add_paragraph("4.2 5G Video MO = Failed")
    elif Video_Orig5G == 'Pass':
        doc.add_paragraph("4.2 5G Video MO = Pass")
    if Video_Term5G == 'Fail':
        doc.add_paragraph("4.2 5G Video MT = Failed")
    elif Video_Term5G == 'Pass':
        doc.add_paragraph("4.2 5G Video MT = Pass")
    if Video_Lc5G == 'Fail':
        doc.add_paragraph("4.3 5G Video LC = Failed")
    elif Video_Lc5G == 'Pass':
        doc.add_paragraph("4.3 5G Video LC = Pass")

def Tab_CAT_Sum(ExeCAT1Summ):
    #CAT-M1 Features
    CATM1_GSMA = ExeCAT1Summ['C9'].value
    CATM1_Perform = ExeCAT1Summ['C10'].value
    CATM1_UICC = ExeCAT1Summ['C11'].value
    CATM1_SMS = ExeCAT1Summ['C12'].value
    CATM1_Pwr = ExeCAT1Summ['C13'].value
    if CATM1_GSMA == 'Fail':
        doc.add_paragraph("CAT-M1 Features: GSMA TS.11 = Failed")
    elif CATM1_GSMA == 'Pass':
        doc.add_paragraph("CAT-M1 Features: GSMA TS.11 = Pass")
    if CATM1_Perform == 'Fail':
        doc.add_paragraph("CAT-M1 Features: Performance = Failed")
    elif CATM1_Perform == 'Pass':
        doc.add_paragraph("CAT-M1 Features: Performance = Pass")
    if CATM1_UICC == 'Fail':
        doc.add_paragraph("CAT-M1 Features: UICC = Failed")
    elif CATM1_UICC == 'Pass':
        doc.add_paragraph("CAT-M1 Features: UICC = Pass")
    if CATM1_SMS == 'Fail':
        doc.add_paragraph("CAT-M1 Features: SMS over NAS = Failed")
    elif CATM1_SMS == 'Pass':
        doc.add_paragraph("CAT-M1 Features: SMS over NAS = Pass")
    if CATM1_Pwr == 'Fail':
        doc.add_paragraph("CAT-M1 Features: Power Savings Techniques = Failed")
    elif CATM1_Pwr == 'Pass':
        doc.add_paragraph("CAT-M1 Features: Power Savings Techniques = Pass")

    #Data Performance
    CATM1_DL = ExeCAT1Summ['B27'].value   #2.14
    CATM1_UL = ExeCAT1Summ['B28'].value   #2.15
    #Mixed LTE Calls
    CATM1_Origs = ExeCAT1Summ['C32'].value  # 2.18
    CATM1_Terms = ExeCAT1Summ['C33'].value  # 2.19
    CATM1_LC = ExeCAT1Summ['C34'].value  # 2.20
    if CATM1_DL == 'Fail':
        doc.add_paragraph("2.14 CAT-M1 FTP DL = Failed")
    elif CATM1_DL == 'Pass':
        doc.add_paragraph("2.14 CAT-M1 FTP DL = Pass")
    if CATM1_UL == 'Fail':
        doc.add_paragraph("2.15 CAT-M1 FTP UL = Failed")
    elif CATM1_UL == 'Pass':
        doc.add_paragraph("2.15 CAT-M1 FTP UL = Pass")
    if CATM1_Origs == 'Fail':
        doc.add_paragraph("2.18 CAT-M1 Origs (Ping) = Failed")
    elif CATM1_Origs == 'Pass':
        doc.add_paragraph("2.18 CAT-M1 Origs (Ping) = Pass")
    if CATM1_Terms == 'Fail':
        doc.add_paragraph("2.19 CAT-M1 Terms (SMS) = Failed")
    elif CATM1_Terms == 'Pass':
        doc.add_paragraph("2.19 CAT-M1 Terms (SMS) = Pass")
    if CATM1_LC == 'Fail':
        doc.add_paragraph("2.20 CAT-M1 Long Call (Data) = Failed")
    elif CATM1_LC == 'Pass':
        doc.add_paragraph("2.20 CAT-M1 Long Call (Data) = Pass")

def Tab_NB_Sum(ExeNBSumm):
    #CAT-M1 Features
    NB_GSMA = ExeNBSumm['H9'].value
    NB_Perform = ExeNBSumm['H10'].value
    NB_UICC = ExeNBSumm['H11'].value
    NB_SMS = ExeNBSumm['H12'].value
    NB_Pwr = ExeNBSumm['H13'].value
    #Data Performance
    NB_DL = ExeNBSumm['G27'].value   #2.14
    NB_UL = ExeNBSumm['G28'].value   #2.15
    #Mixed LTE Calls
    NB_Origs = ExeNBSumm['H32'].value  # 2.18
    NB_Terms = ExeNBSumm['H33'].value  # 2.19
    NB_LC = ExeNBSumm['H34'].value  # 2.20
    if NB_GSMA == 'Fail':
        doc.add_paragraph("NB Features: GSMA TS.11 = Failed")
    elif NB_GSMA == 'Pass':
        doc.add_paragraph("NB Features: GSMA TS.11 = Pass")
    if NB_Perform == 'Fail':
        doc.add_paragraph("NB Features: Performance = Failed")
    elif NB_Perform == 'Pass':
        doc.add_paragraph("NB Features: Performance = Pass")
    if NB_UICC == 'Fail':
        doc.add_paragraph("NB Features: UICC = Failed")
    elif NB_UICC == 'Pass':
        doc.add_paragraph("NB Features: UICC = Pass")
    if NB_SMS == 'Fail':
        doc.add_paragraph("NB Features: SMS over NAS = Failed")
    elif NB_SMS == 'Pass':
        doc.add_paragraph("NB Features: SMS over NAS = Pass")
    if NB_Pwr == 'Fail':
        doc.add_paragraph("NB Features: Power Savings Techniques = Failed")
    elif NB_Pwr == 'Pass':
        doc.add_paragraph("NB Features: Power Savings Techniques = Pass")
    if NB_DL == 'Fail':
        doc.add_paragraph("2.14 NB FTP DL = Failed")
    elif NB_DL == 'Pass':
        doc.add_paragraph("2.14 NB FTP DL = Pass")
    if NB_UL == 'Fail':
        doc.add_paragraph("2.15 NB FTP UL = Failed")
    elif NB_UL == 'Pass':
        doc.add_paragraph("2.15 NB FTP UL = Pass")
    if NB_Origs == 'Fail':
        doc.add_paragraph("2.18 NB Origs (Ping) = Failed")
    elif NB_Origs == 'Pass':
        doc.add_paragraph("2.18 NB Origs (Ping) = Pass")
    if NB_Terms == 'Fail':
        doc.add_paragraph("2.19 NB Terms (SMS) = Failed")
    elif NB_Terms == 'Pass':
        doc.add_paragraph("2.19 NB Terms (SMS) = Pass")
    if NB_LC == 'Fail':
        doc.add_paragraph("2.20 NB Long Call (Data) = Failed")
    elif NB_LC == 'Pass':
        doc.add_paragraph("2.20 NB Long Call (Data) = Pass")

Exe4GSumm2 = wb.worksheets[3]
Tab_4G_Sum(Exe4GSumm2)
Exe5GSumm2 = wb.worksheets[4]
Tab_5G_Sum(Exe5GSumm2)
ExeCAT1Summ2 = wb.worksheets[5]
Tab_CAT_Sum(ExeCAT1Summ2)
Tab_NB_Sum(ExeCAT1Summ2)
doc.save('test.docx')